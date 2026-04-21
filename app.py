from __future__ import annotations

import os
import math
import tkinter as tk
from tkinter import ttk, messagebox, filedialog

from core.line_rate.conductor_loader import load_conductor_database, ConductorDatabase
from core.line_rate.ieee738 import calculate_steady_state_rating
from core.line_rate.solar_ieee738 import parse_date_input, parse_time_input
from core.line_rate.conductor import Conductor
from gui.tlic.main_form import MainForm as TlicMainForm


FT_PER_M = 3.280839895013123


class LineRatingApp(tk.Tk):
    def __init__(self) -> None:
        super().__init__()

        self.title("Line Rate / TLIC")
        self.geometry("1440x900")
        self.minsize(1180, 720)

        self.database: ConductorDatabase | None = None
        self.selected_conductor: Conductor | None = None

        self.family_var = tk.StringVar()
        self.conductor_var = tk.StringVar()

        self._build_ui()
        self._load_database()

    def _build_ui(self) -> None:
        self.ribbon = ttk.Notebook(self)
        self.ribbon.pack(fill="both", expand=True)

        self.line_rate_tab = ttk.Frame(self.ribbon)
        self.ribbon.add(self.line_rate_tab, text="Line Rate")

        self.tlic_tab = TlicMainForm(self.ribbon, menu_parent=self, build_menu=False)
        self.ribbon.add(self.tlic_tab, text="TLIC")

        top_frame = ttk.Frame(self.line_rate_tab, padding=10)
        top_frame.pack(fill="x")

        ttk.Label(top_frame, text="Conductor Family:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        self.family_combo = ttk.Combobox(top_frame, textvariable=self.family_var, state="readonly", width=20)
        self.family_combo.grid(row=0, column=1, sticky="w", padx=(0, 20))
        self.family_combo.bind("<<ComboboxSelected>>", self._on_family_changed)

        ttk.Label(top_frame, text="Conductor:").grid(row=0, column=2, sticky="w", padx=(0, 8))
        self.conductor_combo = ttk.Combobox(top_frame, textvariable=self.conductor_var, state="readonly", width=30)
        self.conductor_combo.grid(row=0, column=3, sticky="w", padx=(0, 20))
        self.conductor_combo.bind("<<ComboboxSelected>>", self._on_conductor_changed)

        self.reload_button = ttk.Button(top_frame, text="Reload Data", command=self._load_database)
        self.reload_button.grid(row=0, column=4, sticky="w")

        middle_frame = ttk.Frame(self.line_rate_tab, padding=(10, 0, 10, 10))
        middle_frame.pack(fill="both", expand=True)

        left_frame = ttk.LabelFrame(middle_frame, text="Conductor Data", padding=10)
        left_frame.pack(side="left", fill="both", expand=True, padx=(0, 5))

        self.data_tree = ttk.Treeview(
            left_frame,
            columns=("property", "value"),
            show="headings",
            height=18
        )
        self.data_tree.heading("property", text="Property")
        self.data_tree.heading("value", text="Value")
        self.data_tree.column("property", width=290, anchor="w")
        self.data_tree.column("value", width=250, anchor="w")
        self.data_tree.pack(fill="both", expand=True)

        right_frame = ttk.LabelFrame(middle_frame, text="Rating Inputs", padding=10)
        right_frame.pack(side="right", fill="y", padx=(5, 0))

        # These defaults match the saved workbook example in Resources/LineRatingIEEE.xlsx.
        self.input_vars = {
            "ambient_temp_c": tk.StringVar(value="40"),
            "wind_speed_mps": tk.StringVar(value="2"),        # entered as ft/s, converted to m/s
            "wind_angle_deg": tk.StringVar(value="45"),       # beta = 0 in workbook
            "elevation_m": tk.StringVar(value="500"),         # entered as ft, converted to m
            "target_temp_c": tk.StringVar(value="100"),
            "emissivity": tk.StringVar(value="0.91"),          # Original Value From LineRate Tool
            "absorptivity": tk.StringVar(value="0.91"),        # Original Value From LineRate Tool
            "latitude_deg": tk.StringVar(value="32.234"),      # 32.234 is the given latitude for the central region of DESC
            "line_azimuth_deg": tk.StringVar(value="0"),
            "line_voltage_kv": tk.StringVar(value="115"),
            "date_text": tk.StringVar(value="7/1/2026"),
            "time_text": tk.StringVar(value="12:00 PM"),
            "atmosphere_type": tk.StringVar(value="clear"),
            "r25_override": tk.StringVar(value=""),
            "r75_override": tk.StringVar(value=""),
            "r200_override": tk.StringVar(value=""),
        }

        input_rows = [
            ("Ambient Temp (C)", "ambient_temp_c"),
            ("Wind Velocity (ft/s)", "wind_speed_mps"),
            ("Wind Angle to Axis (deg)", "wind_angle_deg"),
            ("Altitude (ft)", "elevation_m"),
            ("Max MOT (C)", "target_temp_c"),
            ("Emissivity", "emissivity"),
            ("Absorptivity", "absorptivity"),
            ("Latitude (deg)", "latitude_deg"),
            ("Line Azimuth (deg)", "line_azimuth_deg"),
            ("System Voltage (kV L-L)", "line_voltage_kv"),
            ("Date", "date_text"),
            ("Time", "time_text"),
            ("R25 Override (ohm/mi)", "r25_override"),
            ("R75 Override (ohm/mi)", "r75_override"),
            ("R200 Override (ohm/mi)", "r200_override"),
        ]

        for row_idx, (label, key) in enumerate(input_rows):
            ttk.Label(right_frame, text=label).grid(row=row_idx, column=0, sticky="w", pady=4, padx=(0, 8))
            if key == "line_voltage_kv":
                voltage_combo = ttk.Combobox(
                    right_frame,
                    textvariable=self.input_vars[key],
                    state="readonly",
                    values=["33", "44", "115", "230"],
                    width=15,
                )
                voltage_combo.grid(row=row_idx, column=1, sticky="w", pady=4)
            else:
                ttk.Entry(right_frame, textvariable=self.input_vars[key], width=18).grid(
                    row=row_idx, column=1, sticky="w", pady=4
                )

        atmosphere_row = len(input_rows)
        ttk.Label(right_frame, text="Atmosphere").grid(row=atmosphere_row, column=0, sticky="w", pady=4, padx=(0, 8))
        atmosphere_combo = ttk.Combobox(
            right_frame,
            textvariable=self.input_vars["atmosphere_type"],
            state="readonly",
            values=["clear", "industrial"],
            width=15
        )
        atmosphere_combo.grid(row=atmosphere_row, column=1, sticky="w", pady=4)

        ttk.Separator(right_frame, orient="horizontal").grid(
            row=atmosphere_row + 1, column=0, columnspan=2, sticky="ew", pady=12
        )

        self.calculate_button = ttk.Button(right_frame, text="Calculate Rating", command=self._calculate_rating)
        self.calculate_button.grid(row=atmosphere_row + 2, column=0, columnspan=2, sticky="ew")

        self.show_math_button = ttk.Button(right_frame, text="Show Math", command=self._show_math_window)
        self.show_math_button.grid(row=atmosphere_row + 3, column=0, columnspan=2, sticky="ew", pady=(8, 0))

        result_frame = ttk.LabelFrame(self.line_rate_tab, text="Calculation Result", padding=10)
        result_frame.pack(fill="both", expand=False, padx=10, pady=(0, 10))

        self.result_text = tk.Text(result_frame, height=14, wrap="word")
        self.result_text.pack(fill="both", expand=True)
        self.result_text.insert("1.0", "Results will appear here after calculation.\n")
        self.result_text.configure(state="disabled")

        self.status_var = tk.StringVar(value="Ready.")
        status_bar = ttk.Label(self, textvariable=self.status_var, relief="sunken", anchor="w", padding=6)
        status_bar.pack(fill="x", side="bottom")

    def _find_data_source(self) -> str:
        resources_dir = os.path.join(os.path.dirname(__file__), "Resources")
        preferred_files = [
            "ConData.xlsx",
            "ConductorData.xlsx",
        ]

        for filename in preferred_files:
            path = os.path.join(resources_dir, filename)
            if os.path.exists(path):
                return path

        raise FileNotFoundError("No conductor workbook found in Resources. Expected ConData.xlsx or ConductorData.xlsx.")

    def _load_database(self) -> None:
        try:
            filepath = self._find_data_source()
            self.database = load_conductor_database(filepath)

            families = self.database.get_families()
            self.family_combo["values"] = families

            if not families:
                self.family_var.set("")
                self.conductor_combo["values"] = []
                self._clear_data_tree()
                self.status_var.set("No conductor sheets found.")
                return

            # Prefer ACSR if present since that matches your current boss case.
            first_family = "ACSR" if "ACSR" in families else families[0]
            self.family_var.set(first_family)
            self._populate_conductors(first_family)

            total_count = sum(len(self.database.get_conductors(f)) for f in families)
            source_name = os.path.basename(filepath)
            self.status_var.set(
                f"Loaded {total_count} conductors from {len(families)} family/families using {source_name}."
            )

        except Exception as exc:
            messagebox.showerror("Load Error", str(exc))
            self.status_var.set("Failed to load conductor database.")

    def _populate_conductors(self, family: str) -> None:
        if self.database is None:
            return

        conductors = self.database.get_conductors(family)
        names = [c.code_word for c in conductors if c.code_word]

        self.conductor_combo["values"] = names

        if names:
            preferred = "BITTERN" if "BITTERN" in names else names[0]
            self.conductor_var.set(preferred)
            self._display_selected_conductor(family, preferred)
        else:
            self.conductor_var.set("")
            self.selected_conductor = None
            self._clear_data_tree()
            self.status_var.set(f"No conductors found in family '{family}'.")

    def _on_family_changed(self, _event=None) -> None:
        family = self.family_var.get().strip()
        self._populate_conductors(family)

    def _on_conductor_changed(self, _event=None) -> None:
        family = self.family_var.get().strip()
        code_word = self.conductor_var.get().strip()
        self._display_selected_conductor(family, code_word)

    def _display_selected_conductor(self, family: str, code_word: str) -> None:
        if self.database is None:
            return

        conductor = self.database.find_conductor(family, code_word)
        self.selected_conductor = conductor
        self._clear_data_tree()

        if conductor is None:
            self.status_var.set(f"Could not find conductor '{code_word}' in '{family}'.")
            return

        self.input_vars["target_temp_c"].set(self._format_mot_default(conductor))

        data_rows = [
            ("Family", conductor.family),
            ("Code Word", conductor.code_word),
            ("Name", conductor.name),
            ("Size (kcmil)", conductor.size_kcmil),
            ("Stranding", conductor.stranding),
            ("OD (in)", conductor.od_in),
            ("DC Res @20C (ohm/mile)", conductor.dc_res_20c_ohm_per_mile),
            ("AC Res @25C (ohm/mile)", conductor.ac_res_25c_ohm_per_mile),
            ("AC Res @50C (ohm/mile)", conductor.ac_res_50c_ohm_per_mile),
            ("AC Res @75C (ohm/mile)", conductor.ac_res_75c_ohm_per_mile),
            ("AC Res @200C (ohm/mile)", conductor.ac_res_200c_ohm_per_mile),
            ("AC Res @250C (ohm/mile)", conductor.ac_res_250c_ohm_per_mile),
            ("STDOL", conductor.stdol),
            ("GMR (ft)", conductor.gmr_ft),
            ("Xa @60Hz (ohm/mile)", conductor.xa_60hz_ohm_per_mile),
            ("Capacitive Reactance", conductor.capacitive_reactance),
            ("Ampacity / RATEB / STDOL", conductor.ampacity_75c_amp),
        ]

        for prop, value in data_rows:
            self.data_tree.insert("", "end", values=(prop, "" if value is None else value))

        self.status_var.set(f"Selected {family} / {code_word}")

    def _clear_data_tree(self) -> None:
        for item in self.data_tree.get_children():
            self.data_tree.delete(item)

    def _set_result_text(self, text: str) -> None:
        self.result_text.configure(state="normal")
        self.result_text.delete("1.0", "end")
        self.result_text.insert("1.0", text)
        self.result_text.configure(state="disabled")

    @staticmethod
    def _format_optional_float(value: float | None, precision: int = 6) -> str:
        if value is None:
            return ""
        return f"{value:.{precision}f}"

    @staticmethod
    def _fmt(value: float, precision: int = 6) -> str:
        return f"{value:.{precision}f}"

    def _get_float_input(self, key: str, label: str, allow_blank: bool = False):
        raw = self.input_vars[key].get().strip()
        if allow_blank and raw == "":
            return None
        try:
            return float(raw)
        except ValueError:
            raise ValueError(f"Invalid numeric value for {label}: '{raw}'")

    @staticmethod
    def _normalize_conductor_kind(conductor: Conductor) -> str:
        family = (conductor.family or "").strip().upper()
        name = (conductor.name or "").strip().upper()
        code_word = (conductor.code_word or "").strip().upper()
        combined = f"{family} {name} {code_word}"

        if "ACCC" in combined:
            return "ACCC"
        if "ACCR" in combined:
            return "ACCR"
        if "ACSS" in combined:
            return "ACSS"
        if "HYTHERM" in combined or ("CU" in combined and "HY" in combined):
            return "CU-HYTHERM"
        if family == "CU" or " CU " in f" {combined} ":
            return "CU"
        return family or "UNKNOWN"

    def _default_mot_for_conductor(self, conductor: Conductor) -> float:
        kind = self._normalize_conductor_kind(conductor)
        if kind == "CU-HYTHERM":
            return 150.0
        if kind == "ACCC":
            return 200.0
        if kind == "ACCR":
            return 240.0
        if kind == "ACSS":
            return 250.0
        return 100.0

    def _format_mot_default(self, conductor: Conductor) -> str:
        value = self._default_mot_for_conductor(conductor)
        return str(int(value)) if value.is_integer() else str(value)

    def _rating_targets_for_conductor(self, conductor: Conductor, requested_max_temp_c: float) -> dict[str, float]:
        kind = self._normalize_conductor_kind(conductor)

        if kind == "CU-HYTHERM":
            max_temp = min(requested_max_temp_c + 25.0, 150.0)
            return {
                "A": min(max_temp, 119.0),
                "B": min(max_temp, 125.0),
                "C": max_temp,
            }
        if kind == "ACCC":
            max_temp = min(requested_max_temp_c, 200.0)
            return {
                "A": min(max_temp, 180.0),
                "B": min(max_temp, 200.0),
                "C": max_temp,
            }
        if kind == "ACCR":
            max_temp = min(requested_max_temp_c, 240.0)
            return {
                "A": min(max_temp, 210.0),
                "B": min(max_temp, 240.0),
                "C": max_temp,
            }
        if kind == "ACSS":
            max_temp = min(requested_max_temp_c, 250.0)
            return {
                "A": min(max_temp, 200.0),
                "B": min(max_temp, 250.0),
                "C": max_temp,
            }

        # Legacy default used for ACSR and copper.
        return {
            "A": min(requested_max_temp_c, 94.0),
            "B": min(requested_max_temp_c, 100.0),
            "C": requested_max_temp_c,
        }

    def _calculate_payload(self) -> dict:
        if self.selected_conductor is None:
            raise ValueError("Please select a conductor first.")

        ambient_temp_c = self._get_float_input("ambient_temp_c", "Ambient Temp")
        wind_speed_fps = self._get_float_input("wind_speed_mps", "Wind Velocity")
        wind_speed_mps = wind_speed_fps / FT_PER_M
        wind_angle_deg = self._get_float_input("wind_angle_deg", "Wind Angle")
        altitude_ft = self._get_float_input("elevation_m", "Altitude")
        elevation_m = altitude_ft / FT_PER_M
        target_temp_c = self._get_float_input("target_temp_c", "Max MOT")
        emissivity = self._get_float_input("emissivity", "Emissivity")
        absorptivity = self._get_float_input("absorptivity", "Absorptivity")
        latitude_deg = self._get_float_input("latitude_deg", "Latitude")
        line_azimuth_deg = self._get_float_input("line_azimuth_deg", "Line Azimuth")
        line_voltage_kv = self._get_float_input("line_voltage_kv", "System Voltage", allow_blank=True)
        r25_override = self._get_float_input("r25_override", "R25 Override", allow_blank=True)
        r75_override = self._get_float_input("r75_override", "R75 Override", allow_blank=True)
        r200_override = self._get_float_input("r200_override", "R200 Override", allow_blank=True)

        input_date = parse_date_input(self.input_vars["date_text"].get())
        input_time = parse_time_input(self.input_vars["time_text"].get())
        atmosphere_type = self.input_vars["atmosphere_type"].get().strip().lower()

        rating_targets = self._rating_targets_for_conductor(self.selected_conductor, target_temp_c)
        rating_results: dict[str, dict] = {}
        for rating_name, rating_temp_c in rating_targets.items():
            rating_results[rating_name] = calculate_steady_state_rating(
                conductor=self.selected_conductor,
                ambient_temp_c=ambient_temp_c,
                wind_speed_mps=wind_speed_mps,
                wind_angle_deg=wind_angle_deg,
                elevation_m=elevation_m,
                target_temp_c=rating_temp_c,
                emissivity=emissivity,
                absorptivity=absorptivity,
                latitude_deg=latitude_deg,
                line_azimuth_deg=line_azimuth_deg,
                input_date=input_date,
                input_time=input_time,
                atmosphere_type=atmosphere_type,
                r25_override=r25_override,
                r75_override=r75_override,
                r200_override=r200_override,
            )

        rating_mvas = {
            name: (None if line_voltage_kv is None else (math.sqrt(3.0) * line_voltage_kv * rating["amps"] / 1000.0))
            for name, rating in rating_results.items()
        }

        return {
            "ambient_temp_c": ambient_temp_c,
            "wind_speed_mps": wind_speed_mps,
            "wind_speed_fps": wind_speed_fps,
            "wind_angle_deg": wind_angle_deg,
            "elevation_m": elevation_m,
            "altitude_ft": altitude_ft,
            "target_temp_c": target_temp_c,
            "emissivity": emissivity,
            "absorptivity": absorptivity,
            "latitude_deg": latitude_deg,
            "line_azimuth_deg": line_azimuth_deg,
            "line_voltage_kv": line_voltage_kv,
            "r25_override": r25_override,
            "r75_override": r75_override,
            "r200_override": r200_override,
            "input_date": input_date,
            "input_time": input_time,
            "atmosphere_type": atmosphere_type,
            "rating_targets": rating_targets,
            "rating_results": rating_results,
            "rating_mvas": rating_mvas,
            "conductor_kind": self._normalize_conductor_kind(self.selected_conductor),
        }

    def _insert_math_line(self, text_widget: tk.Text, text: str = "", tags: tuple[str, ...] = ()) -> None:
        text_widget.insert("end", text + "\n", tags)

    @staticmethod
    def _bullet(text: str) -> str:
        return f"\u2022\t{text}"

    def _populate_math_text(self, text_widget: tk.Text, payload: dict) -> None:
        conductor = self.selected_conductor
        assert conductor is not None

        text_widget.configure(state="normal")
        text_widget.delete("1.0", "end")

        self._insert_math_line(text_widget, "LINE RATING ENGINEERING CHECK", ("title",))
        self._insert_math_line(
            text_widget,
            f"{conductor.code_word}   |   {conductor.family}   |   {payload['conductor_kind']}",
            ("meta",),
        )
        self._insert_math_line(text_widget)

        self._insert_math_line(text_widget, "Answers", ("section",))
        for rating_name in ("A", "B", "C"):
            amps = payload["rating_results"][rating_name]["amps"]
            mva = payload["rating_mvas"][rating_name]
            self._insert_math_line(
                text_widget,
                self._bullet(f"Rating {rating_name}: {self._fmt(amps)} A | {self._format_optional_float(mva, 6)} MVA"),
                ("body",),
            )
        self._insert_math_line(text_widget)

        self._insert_math_line(text_widget, "Given", ("section",))
        given_lines = [
            f"Source workbook: {os.path.basename(self.database.source_path) if self.database and self.database.source_path else 'unknown'}",
            f"System voltage: {self._format_optional_float(payload['line_voltage_kv'], 3)} kV line-to-line",
            f"Ambient temperature, Ta: {self._fmt(payload['ambient_temp_c'], 3)} C",
            f"Wind velocity, Vw: {self._fmt(payload['wind_speed_fps'], 6)} ft/s",
            f"Wind angle to conductor axis, phi: {self._fmt(payload['wind_angle_deg'], 3)} deg",
            f"Altitude, He: {self._fmt(payload['altitude_ft'], 3)} ft",
            f"Emissivity, eps: {self._fmt(payload['emissivity'], 3)}",
            f"Absorptivity, alpha: {self._fmt(payload['absorptivity'], 3)}",
            f"Latitude: {self._fmt(payload['latitude_deg'], 3)} deg",
            f"Line azimuth, Zl: {self._fmt(payload['line_azimuth_deg'], 3)} deg",
            f"Date: {payload['input_date'].isoformat()}",
            f"Time: {payload['input_time'].strftime('%H:%M:%S')}",
            f"Atmosphere: {payload['atmosphere_type']}",
            f"Max MOT: {self._fmt(payload['target_temp_c'], 3)} C",
        ]
        for line in given_lines:
            self._insert_math_line(text_widget, self._bullet(line), ("body",))
        self._insert_math_line(text_widget)

        self._insert_math_line(text_widget, "Reference Equations", ("section",))
        equations = [
            "I = sqrt((qc + qr - qs) / R(Tavg))",
            "qc = max(qcn, qc1, qc2)",
            "qcn = 1.825 * rho_f^0.5 * D0^0.75 * (Ts - Ta)^1.25",
            "qc1 = Kangle * (1.01 + 1.35 * NRe^0.52) * kf * (Ts - Ta)",
            "qc2 = Kangle * 0.754 * NRe^0.6 * kf * (Ts - Ta)",
            "qr = 1.656 * D0 * eps * [((Ts+273)/100)^4 - ((Ta+273)/100)^4]",
            "qs = alpha * Qse * sin(theta) * D0",
            "MVA = sqrt(3) * kV * A / 1000",
        ]
        for line in equations:
            self._insert_math_line(text_widget, self._bullet(line), ("equation",))
        self._insert_math_line(text_widget)

        for rating_name in ("A", "B", "C"):
            result = payload["rating_results"][rating_name]
            solar = result["solar"]
            target = payload["rating_targets"][rating_name]
            mva = payload["rating_mvas"][rating_name]

            self._insert_math_line(text_widget, f"Rating {rating_name} Calculations", ("section",))

            self._insert_math_line(text_widget, "Resistance", ("subsection",))
            self._insert_math_line(text_widget, self._bullet(f"Mode: {result['resistance_mode']}"), ("body",))
            self._insert_math_line(text_widget, self._bullet(f"R25 = {self._format_optional_float(result['r25_used_ohm_per_mile'])} ohm/mi"), ("body",))
            self._insert_math_line(text_widget, self._bullet(f"R75 = {self._format_optional_float(result['r75_used_ohm_per_mile'])} ohm/mi"), ("body",))
            self._insert_math_line(text_widget, self._bullet(f"R200 = {self._format_optional_float(result['r200_used_ohm_per_mile'])} ohm/mi"), ("body",))
            self._insert_math_line(text_widget, self._bullet(f"R250 = {self._format_optional_float(result['r250_used_ohm_per_mile'])} ohm/mi"), ("body",))
            self._insert_math_line(text_widget, self._bullet(f"R(Tavg) = {self._fmt(result['resistance_ohm_per_mile'])} ohm/mi = {self._fmt(result['resistance_ohm_per_ft'], 9)} ohm/ft"), ("body",))
            self._insert_math_line(text_widget)

            self._insert_math_line(text_widget, "Air Properties And Geometry", ("subsection",))
            self._insert_math_line(text_widget, self._bullet(f"Ts = {self._fmt(target, 3)} C"), ("body",))
            self._insert_math_line(text_widget, self._bullet(f"Ta = {self._fmt(payload['ambient_temp_c'], 3)} C"), ("body",))
            self._insert_math_line(
                text_widget,
                self._bullet(
                    f"Tfilm = (Ts + Ta) / 2 = ({self._fmt(target, 3)} + {self._fmt(payload['ambient_temp_c'], 3)}) / 2 = {self._fmt(result['tfilm_c'])} C"
                ),
                ("equation",),
            )
            self._insert_math_line(
                text_widget,
                self._bullet(f"D0 = {self._fmt(result['diameter_in'], 6)} in = {self._fmt(result['diameter_ft'], 9)} ft"),
                ("body",),
            )
            self._insert_math_line(
                text_widget,
                self._bullet(
                    f"He = {self._fmt(payload['altitude_ft'], 3)} ft = {self._fmt(payload['elevation_m'], 6)} m"
                ),
                ("body",),
            )
            self._insert_math_line(
                text_widget,
                self._bullet(
                    f"rho_f = (0.080695 - 2.901e-6*He + 3.7e-11*He^2) / (1 + 0.00367*Tfilm) = {self._fmt(result['rho_f_lb_per_ft3'], 8)} lb/ft^3"
                ),
                ("equation",),
            )
            self._insert_math_line(
                text_widget,
                self._bullet(
                    f"mu_f = 9.806e-7 * (Tfilm + 273)^1.5 / (Tfilm + 383.4) = {self._fmt(result['mu_f_lb_per_ft_s'], 10)} lb/ft-s"
                ),
                ("equation",),
            )
            self._insert_math_line(
                text_widget,
                self._bullet(
                    f"kf = 0.007388 + 2.279e-5*Tfilm - 1.343e-9*Tfilm^2 = {self._fmt(result['k_f_w_per_ft_c'], 8)} W/ft-C"
                ),
                ("equation",),
            )
            self._insert_math_line(
                text_widget,
                self._bullet(f"Wind speed = {self._fmt(result['wind_speed_fps'], 6)} ft/s"),
                ("body",),
            )
            self._insert_math_line(
                text_widget,
                self._bullet(f"beta = 90 - phi = 90 - {self._fmt(payload['wind_angle_deg'], 3)} = {self._fmt(result['beta_deg'], 3)} deg"),
                ("equation",),
            )
            self._insert_math_line(
                text_widget,
                self._bullet(
                    f"Kangle = 1.194 - sin(beta) - 0.194*cos(2*beta) + 0.368*sin(2*beta) = {self._fmt(result['k_angle'])}"
                ),
                ("equation",),
            )
            self._insert_math_line(
                text_widget,
                self._bullet(
                    f"NRe = D0 * rho_f * Vw / mu_f = {self._fmt(result['diameter_ft'], 9)} * {self._fmt(result['rho_f_lb_per_ft3'], 8)} * {self._fmt(result['wind_speed_fps'], 6)} / {self._fmt(result['mu_f_lb_per_ft_s'], 10)} = {self._fmt(result['n_re'])}"
                ),
                ("equation",),
            )
            self._insert_math_line(text_widget)

            self._insert_math_line(text_widget, "Heat Terms", ("subsection",))
            self._insert_math_line(
                text_widget,
                self._bullet(
                    f"qcn = 1.825 * rho_f^0.5 * D0^0.75 * (Ts - Ta)^1.25 = {self._fmt(result['qcn_w_per_ft'])} W/ft"
                ),
                ("equation",),
            )
            self._insert_math_line(
                text_widget,
                self._bullet(
                    f"qc1 = Kangle * (1.01 + 1.35*NRe^0.52) * kf * (Ts - Ta) = {self._fmt(result['qc1_w_per_ft'])} W/ft"
                ),
                ("equation",),
            )
            self._insert_math_line(
                text_widget,
                self._bullet(
                    f"qc2 = Kangle * 0.754 * NRe^0.6 * kf * (Ts - Ta) = {self._fmt(result['qc2_w_per_ft'])} W/ft"
                ),
                ("equation",),
            )
            self._insert_math_line(
                text_widget,
                self._bullet(f"qc = max(qcn, qc1, qc2) = {self._fmt(result['qc_w_per_ft'])} W/ft"),
                ("equation",),
            )
            self._insert_math_line(
                text_widget,
                self._bullet(
                    f"qr = 1.656 * D0 * eps * [((Ts+273)/100)^4 - ((Ta+273)/100)^4] = {self._fmt(result['qr_w_per_ft'])} W/ft"
                ),
                ("equation",),
            )
            self._insert_math_line(
                text_widget,
                self._bullet(f"qs = {self._fmt(result['qs_w_per_ft'])} W/ft"),
                ("equation",),
            )
            self._insert_math_line(text_widget)

            self._insert_math_line(text_widget, "Solar Geometry", ("subsection",))
            self._insert_math_line(text_widget, self._bullet(f"N = {solar['n_day']}"), ("body",))
            self._insert_math_line(
                text_widget,
                self._bullet(f"omega = (hour - 12) * 15 = {self._fmt(solar['omega_deg'])} deg"),
                ("equation",),
            )
            self._insert_math_line(
                text_widget,
                self._bullet(f"delta = 23.45 * sin(((284 + N) / 365) * 360) = {self._fmt(solar['delta_deg'])} deg"),
                ("equation",),
            )
            self._insert_math_line(
                text_widget,
                self._bullet(
                    f"Hc = asin(cos(Lat)*cos(delta)*cos(omega) + sin(Lat)*sin(delta)) = {self._fmt(solar['hc_deg'])} deg"
                ),
                ("equation",),
            )
            self._insert_math_line(text_widget, self._bullet(f"chi = {self._fmt(solar['chi'])}"), ("body",))
            self._insert_math_line(text_widget, self._bullet(f"C = {self._fmt(solar['c_constant'])} deg"), ("body",))
            self._insert_math_line(text_widget, self._bullet(f"Zc = C + atan(chi) = {self._fmt(solar['zc_deg'])} deg"), ("equation",))
            self._insert_math_line(text_widget, self._bullet(f"theta = acos(cos(Hc) * cos(Zc - Zl)) = {self._fmt(solar['theta_deg'])} deg"), ("equation",))
            self._insert_math_line(text_widget, self._bullet(f"Qs = {self._fmt(solar['qs_sea_level_w_per_ft2'])} W/ft^2"), ("body",))
            self._insert_math_line(text_widget, self._bullet(f"Ksolar = {self._fmt(solar['ksolar'])}"), ("body",))
            self._insert_math_line(text_widget, self._bullet(f"Qse = Ksolar * Qs = {self._fmt(solar['qse_w_per_ft2'])} W/ft^2"), ("equation",))
            self._insert_math_line(text_widget)

            self._insert_math_line(text_widget, "Final Calculation", ("subsection",))
            self._insert_math_line(
                text_widget,
                f"I = sqrt((qc + qr - qs) / R) = sqrt(({self._fmt(result['qc_w_per_ft'])} + {self._fmt(result['qr_w_per_ft'])} - {self._fmt(result['qs_w_per_ft'])}) / {self._fmt(result['resistance_ohm_per_ft'], 9)})",
                ("equation",),
            )
            self._insert_math_line(text_widget, f"I = {self._fmt(result['amps'])} A", ("answer",))
            self._insert_math_line(text_widget, f"3-phase MVA = {self._format_optional_float(mva, 6)}", ("answer",))
            self._insert_math_line(text_widget)
            self._insert_math_line(text_widget, f"Rating {rating_name} Worksheet Final Answer", ("section",))
            self._insert_math_line(
                text_widget,
                f"MOT = {self._fmt(target, 3)} C   |   {self._fmt(result['amps'])} A   |   {self._format_optional_float(mva, 6)} MVA",
                ("highlight",),
            )
            self._insert_math_line(text_widget)

        text_widget.configure(state="disabled")

    def _show_math_window(self) -> None:
        try:
            payload = self._calculate_payload()
        except Exception as exc:
            messagebox.showerror("Math View Error", str(exc))
            self.status_var.set("Math view failed.")
            return

        window = tk.Toplevel(self)
        window.title("Line Rating Math")
        window.geometry("980x760")
        window.minsize(820, 620)

        frame = ttk.Frame(window, padding=10)
        frame.pack(fill="both", expand=True)

        toolbar = ttk.Frame(frame)
        toolbar.grid(row=0, column=0, columnspan=2, sticky="ew", pady=(0, 8))
        ttk.Button(
            toolbar,
            text="Export To Word",
            command=lambda: self._export_math_to_word(payload),
        ).pack(side="left")

        text = tk.Text(
            frame,
            wrap="none",
            font=("Consolas", 10),
            bg="#f6f1e8",
            fg="#18212b",
            insertbackground="#18212b",
            padx=20,
            pady=18,
        )
        yscroll = ttk.Scrollbar(frame, orient="vertical", command=text.yview)
        xscroll = ttk.Scrollbar(frame, orient="horizontal", command=text.xview)
        text.configure(yscrollcommand=yscroll.set, xscrollcommand=xscroll.set)

        text.grid(row=1, column=0, sticky="nsew")
        yscroll.grid(row=1, column=1, sticky="ns")
        xscroll.grid(row=2, column=0, sticky="ew")
        frame.columnconfigure(0, weight=1)
        frame.rowconfigure(1, weight=1)

        text.tag_configure("title", font=("Segoe UI", 18, "bold"), foreground="#102a43", spacing3=8)
        text.tag_configure("meta", font=("Segoe UI", 10, "italic"), foreground="#486581", spacing3=10)
        text.tag_configure("section", font=("Segoe UI", 13, "bold"), foreground="#7c2d12", spacing1=12, spacing3=6)
        text.tag_configure("subsection", font=("Segoe UI", 11, "bold"), foreground="#243b53", spacing1=8, spacing3=3)
        text.tag_configure("body", font=("Segoe UI", 10), foreground="#18212b", lmargin1=12, lmargin2=34, tabs=("30",))
        text.tag_configure("equation", font=("Segoe UI", 10, "bold"), foreground="#1f5f8b", lmargin1=12, lmargin2=34, tabs=("30",), spacing3=2)
        text.tag_configure("answer", font=("Segoe UI", 11, "bold"), foreground="#14532d", lmargin1=12, lmargin2=12, spacing3=2)
        text.tag_configure("highlight", font=("Segoe UI", 11, "bold"), foreground="#111827", background="#fde68a", lmargin1=8, lmargin2=8, spacing3=6)

        self._populate_math_text(text, payload)
        self.status_var.set(f"Opened math view for {self.selected_conductor.code_word}")

    def _export_math_to_word(self, payload: dict) -> None:
        try:
            from docx import Document
        except Exception:
            messagebox.showerror(
                "Export Error",
                "Word export requires the 'python-docx' package to be installed.",
            )
            self.status_var.set("Word export failed.")
            return

        conductor = self.selected_conductor
        if conductor is None:
            messagebox.showerror("Export Error", "No conductor is selected.")
            self.status_var.set("Word export failed.")
            return

        default_name = f"{conductor.code_word}_line_rating_math.docx"
        save_path = filedialog.asksaveasfilename(
            parent=self,
            title="Export Math To Word",
            defaultextension=".docx",
            initialfile=default_name,
            filetypes=[("Word Document", "*.docx")],
        )
        if not save_path:
            return

        try:
            doc = Document()
            doc.add_heading("Line Rating Engineering Check", level=0)
            doc.add_paragraph(f"{conductor.code_word} | {conductor.family} | {payload['conductor_kind']}")

            doc.add_heading("Answers", level=1)
            for rating_name in ("A", "B", "C"):
                amps = payload["rating_results"][rating_name]["amps"]
                mva = payload["rating_mvas"][rating_name]
                doc.add_paragraph(
                    f"Rating {rating_name}: {self._fmt(amps)} A | {self._format_optional_float(mva, 6)} MVA",
                    style="List Bullet",
                )

            doc.add_heading("Given", level=1)
            given_lines = [
                f"Source workbook: {os.path.basename(self.database.source_path) if self.database and self.database.source_path else 'unknown'}",
                f"System voltage: {self._format_optional_float(payload['line_voltage_kv'], 3)} kV line-to-line",
                f"Ambient temperature, Ta: {self._fmt(payload['ambient_temp_c'], 3)} C",
                f"Wind velocity, Vw: {self._fmt(payload['wind_speed_fps'], 6)} ft/s",
                f"Wind angle to conductor axis, phi: {self._fmt(payload['wind_angle_deg'], 3)} deg",
                f"Altitude, He: {self._fmt(payload['altitude_ft'], 3)} ft",
                f"Emissivity, eps: {self._fmt(payload['emissivity'], 3)}",
                f"Absorptivity, alpha: {self._fmt(payload['absorptivity'], 3)}",
                f"Latitude: {self._fmt(payload['latitude_deg'], 3)} deg",
                f"Line azimuth, Zl: {self._fmt(payload['line_azimuth_deg'], 3)} deg",
                f"Date: {payload['input_date'].isoformat()}",
                f"Time: {payload['input_time'].strftime('%H:%M:%S')}",
                f"Atmosphere: {payload['atmosphere_type']}",
                f"Max MOT: {self._fmt(payload['target_temp_c'], 3)} C",
            ]
            for line in given_lines:
                doc.add_paragraph(line, style="List Bullet")

            doc.add_heading("Reference Equations", level=1)
            for line in [
                "I = sqrt((qc + qr - qs) / R(Tavg))",
                "qc = max(qcn, qc1, qc2)",
                "qcn = 1.825 * rho_f^0.5 * D0^0.75 * (Ts - Ta)^1.25",
                "qc1 = Kangle * (1.01 + 1.35 * NRe^0.52) * kf * (Ts - Ta)",
                "qc2 = Kangle * 0.754 * NRe^0.6 * kf * (Ts - Ta)",
                "qr = 1.656 * D0 * eps * [((Ts+273)/100)^4 - ((Ta+273)/100)^4]",
                "qs = alpha * Qse * sin(theta) * D0",
                "MVA = sqrt(3) * kV * A / 1000",
            ]:
                doc.add_paragraph(line, style="List Bullet")

            for rating_name in ("A", "B", "C"):
                result = payload["rating_results"][rating_name]
                solar = result["solar"]
                target = payload["rating_targets"][rating_name]
                mva = payload["rating_mvas"][rating_name]

                doc.add_heading(f"Rating {rating_name} Calculations", level=1)

                doc.add_heading("Resistance", level=2)
                for line in [
                    f"Mode: {result['resistance_mode']}",
                    f"R25 = {self._format_optional_float(result['r25_used_ohm_per_mile'])} ohm/mi",
                    f"R75 = {self._format_optional_float(result['r75_used_ohm_per_mile'])} ohm/mi",
                    f"R200 = {self._format_optional_float(result['r200_used_ohm_per_mile'])} ohm/mi",
                    f"R250 = {self._format_optional_float(result['r250_used_ohm_per_mile'])} ohm/mi",
                    f"R(Tavg) = {self._fmt(result['resistance_ohm_per_mile'])} ohm/mi = {self._fmt(result['resistance_ohm_per_ft'], 9)} ohm/ft",
                ]:
                    doc.add_paragraph(line, style="List Bullet")

                doc.add_heading("Air Properties And Geometry", level=2)
                for line in [
                    f"Ts = {self._fmt(target, 3)} C",
                    f"Ta = {self._fmt(payload['ambient_temp_c'], 3)} C",
                    f"Tfilm = {self._fmt(result['tfilm_c'])} C",
                    f"D0 = {self._fmt(result['diameter_in'], 6)} in = {self._fmt(result['diameter_ft'], 9)} ft",
                    f"rho_f = {self._fmt(result['rho_f_lb_per_ft3'], 8)} lb/ft^3",
                    f"mu_f = {self._fmt(result['mu_f_lb_per_ft_s'], 10)} lb/ft-s",
                    f"kf = {self._fmt(result['k_f_w_per_ft_c'], 8)} W/ft-C",
                    f"Wind speed = {self._fmt(result['wind_speed_fps'], 6)} ft/s",
                    f"beta = {self._fmt(result['beta_deg'], 3)} deg",
                    f"Kangle = {self._fmt(result['k_angle'])}",
                    f"NRe = {self._fmt(result['n_re'])}",
                ]:
                    doc.add_paragraph(line, style="List Bullet")

                doc.add_heading("Heat Terms", level=2)
                for line in [
                    f"qcn = {self._fmt(result['qcn_w_per_ft'])} W/ft",
                    f"qc1 = {self._fmt(result['qc1_w_per_ft'])} W/ft",
                    f"qc2 = {self._fmt(result['qc2_w_per_ft'])} W/ft",
                    f"qc = {self._fmt(result['qc_w_per_ft'])} W/ft",
                    f"qr = {self._fmt(result['qr_w_per_ft'])} W/ft",
                    f"qs = {self._fmt(result['qs_w_per_ft'])} W/ft",
                ]:
                    doc.add_paragraph(line, style="List Bullet")

                doc.add_heading("Solar Geometry", level=2)
                for line in [
                    f"N = {solar['n_day']}",
                    f"omega = {self._fmt(solar['omega_deg'])} deg",
                    f"delta = {self._fmt(solar['delta_deg'])} deg",
                    f"Hc = {self._fmt(solar['hc_deg'])} deg",
                    f"chi = {self._fmt(solar['chi'])}",
                    f"C = {self._fmt(solar['c_constant'])} deg",
                    f"Zc = {self._fmt(solar['zc_deg'])} deg",
                    f"theta = {self._fmt(solar['theta_deg'])} deg",
                    f"Qs = {self._fmt(solar['qs_sea_level_w_per_ft2'])} W/ft^2",
                    f"Ksolar = {self._fmt(solar['ksolar'])}",
                    f"Qse = {self._fmt(solar['qse_w_per_ft2'])} W/ft^2",
                ]:
                    doc.add_paragraph(line, style="List Bullet")

                doc.add_heading("Final Calculation", level=2)
                doc.add_paragraph(
                    f"I = sqrt((qc + qr - qs) / R) = sqrt(({self._fmt(result['qc_w_per_ft'])} + {self._fmt(result['qr_w_per_ft'])} - {self._fmt(result['qs_w_per_ft'])}) / {self._fmt(result['resistance_ohm_per_ft'], 9)})"
                )
                doc.add_paragraph(f"I = {self._fmt(result['amps'])} A")
                doc.add_paragraph(f"3-phase MVA = {self._format_optional_float(mva, 6)}")

                doc.add_heading(f"Rating {rating_name} Worksheet Final Answer", level=2)
                doc.add_paragraph(
                    f"MOT = {self._fmt(target, 3)} C | {self._fmt(result['amps'])} A | {self._format_optional_float(mva, 6)} MVA"
                )

            doc.save(save_path)
            self.status_var.set(f"Exported Word file to {os.path.basename(save_path)}")
        except Exception as exc:
            messagebox.showerror("Export Error", str(exc))
            self.status_var.set("Word export failed.")

    def _calculate_rating(self) -> None:
        if self.selected_conductor is None:
            messagebox.showwarning("No Conductor", "Please select a conductor first.")
            return

        try:
            payload = self._calculate_payload()

            ambient_temp_c = payload["ambient_temp_c"]
            wind_speed_mps = payload["wind_speed_mps"]
            wind_angle_deg = payload["wind_angle_deg"]
            altitude_ft = payload["altitude_ft"]
            target_temp_c = payload["target_temp_c"]
            emissivity = payload["emissivity"]
            absorptivity = payload["absorptivity"]
            latitude_deg = payload["latitude_deg"]
            line_azimuth_deg = payload["line_azimuth_deg"]
            line_voltage_kv = payload["line_voltage_kv"]
            input_date = payload["input_date"]
            input_time = payload["input_time"]
            atmosphere_type = payload["atmosphere_type"]
            rating_targets = payload["rating_targets"]
            rating_results = payload["rating_results"]
            result = rating_results["C"]

            source_ref = os.path.basename(self.database.source_path) if self.database and self.database.source_path else "unknown"
            workbook_reference = self.selected_conductor.ampacity_75c_amp
            conductor_kind = payload["conductor_kind"]
            rating_mvas = payload["rating_mvas"]

            result_lines = [
                f"{self.selected_conductor.code_word}  |  {self.selected_conductor.family}  |  {source_ref}",
                "",
                "RATINGS",
                f"  A  |  {rating_results['A']['amps']:.3f} A  |  {self._format_optional_float(rating_mvas['A'], precision=3)} MVA",
                f"  B  |  {rating_results['B']['amps']:.3f} A  |  {self._format_optional_float(rating_mvas['B'], precision=3)} MVA",
                f"  C  |  {rating_results['C']['amps']:.3f} A  |  {self._format_optional_float(rating_mvas['C'], precision=3)} MVA",
                "",
                "TEMPERATURE TARGETS",
                f"  Rule  |  {conductor_kind}",
                f"  A     |  {rating_targets['A']:.3f} C",
                f"  B     |  {rating_targets['B']:.3f} C",
                f"  C     |  {rating_targets['C']:.3f} C",
                "",
                "CASE",
                f"  Ambient  |  {ambient_temp_c:.1f} C",
                f"  Wind     |  {payload['wind_speed_fps']:.3f} ft/s @ {wind_angle_deg:.1f} deg",
                f"  Altitude |  {altitude_ft:.1f} ft",
                f"  Solar    |  {input_date.isoformat()}  {input_time.strftime('%H:%M:%S')}  {atmosphere_type}",
                f"  Voltage  |  {self._format_optional_float(line_voltage_kv, precision=3)} kV",
                "",
                "DETAIL",
                "  Use Show Math for the full equation-by-equation check view.",
            ]

            self._set_result_text("\n".join(result_lines))
            self.status_var.set(f"Calculated steady-state rating for {self.selected_conductor.code_word}")

        except Exception as exc:
            messagebox.showerror("Calculation Error", str(exc))
            self.status_var.set("Calculation failed.")
